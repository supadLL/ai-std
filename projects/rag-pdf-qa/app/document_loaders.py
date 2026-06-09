import csv
from html.parser import HTMLParser
from io import BytesIO, StringIO
from dataclasses import dataclass
from pathlib import Path

from app.ocr_extractor import OcrExtractionError, extract_ocr_text_from_image_bytes


@dataclass(frozen=True)
class ParsedSection:
    section_number: int
    title: str | None
    text: str
    extraction_method: str = "text"


@dataclass(frozen=True)
class ParsedDocument:
    filename: str
    file_type: str
    char_count: int
    preview: str
    sections: list[ParsedSection]


class DocumentLoadError(RuntimeError):
    pass


def load_document_from_bytes(
    filename: str,
    content: bytes,
    preview_chars: int = 500,
    enable_image_ocr: bool = False,
    ocr_language: str = "chi_sim+eng",
) -> ParsedDocument:
    if not content:
        raise DocumentLoadError("Document file is empty")

    suffix = Path(filename).suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        return _load_text_document_from_bytes(filename=filename, content=content, preview_chars=preview_chars)
    if suffix == ".docx":
        return _load_docx_document_from_bytes(
            filename=filename,
            content=content,
            preview_chars=preview_chars,
            enable_image_ocr=enable_image_ocr,
            ocr_language=ocr_language,
        )
    if suffix == ".csv":
        return _load_csv_document_from_bytes(filename=filename, content=content, preview_chars=preview_chars)
    if suffix == ".xlsx":
        return _load_xlsx_document_from_bytes(filename=filename, content=content, preview_chars=preview_chars)
    if suffix in {".html", ".htm"}:
        return _load_html_document_from_bytes(filename=filename, content=content, preview_chars=preview_chars)
    raise DocumentLoadError("Only Markdown, txt, docx, csv, xlsx, html, and htm files are supported by document loaders")


def load_text_document_from_bytes(filename: str, content: bytes, preview_chars: int = 500) -> ParsedDocument:
    return _load_text_document_from_bytes(filename=filename, content=content, preview_chars=preview_chars)


def is_supported_document(filename: str) -> bool:
    suffix = Path(filename).suffix.lower()
    return suffix in {".md", ".markdown", ".txt", ".docx", ".csv", ".xlsx", ".html", ".htm"}


def is_supported_text_document(filename: str) -> bool:
    suffix = Path(filename).suffix.lower()
    return suffix in {".md", ".markdown", ".txt"}


def _load_text_document_from_bytes(filename: str, content: bytes, preview_chars: int = 500) -> ParsedDocument:
    file_type = _detect_text_file_type(filename)
    text = _decode_text(content)
    normalized = text.strip()
    if not normalized:
        raise DocumentLoadError("Document has no text content")

    if file_type == "markdown":
        sections = _parse_markdown_sections(normalized)
    else:
        sections = [ParsedSection(section_number=1, title=None, text=normalized)]

    return ParsedDocument(
        filename=filename,
        file_type=file_type,
        char_count=len(normalized),
        preview=normalized[:preview_chars],
        sections=sections,
    )


def _detect_text_file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".md", ".markdown"}:
        return "markdown"
    if suffix == ".txt":
        return "text"
    raise DocumentLoadError("Only Markdown and txt files are supported by the text document loader")


def _load_docx_document_from_bytes(
    filename: str,
    content: bytes,
    preview_chars: int = 500,
    enable_image_ocr: bool = False,
    ocr_language: str = "chi_sim+eng",
) -> ParsedDocument:
    try:
        from docx import Document

        doc = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentLoadError(f"Failed to load docx document: {exc}") from exc

    sections: list[ParsedSection] = []
    paragraph_parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraph_parts.append(text)

    if paragraph_parts:
        sections.append(
            ParsedSection(
                section_number=len(sections) + 1,
                title="docx paragraphs",
                text="\n".join(paragraph_parts),
                extraction_method="text",
            )
        )

    table_parts: list[str] = []
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            values = [value for value in values if value]
            if values:
                table_parts.append(f"table {table_index} row {row_index}: " + " | ".join(values))

    if table_parts:
        sections.append(
            ParsedSection(
                section_number=len(sections) + 1,
                title="docx tables",
                text="\n".join(table_parts),
                extraction_method="table",
            )
        )

    if enable_image_ocr:
        sections.extend(
            _extract_docx_image_ocr_sections(
                doc=doc,
                filename=filename,
                start_section_number=len(sections) + 1,
                ocr_language=ocr_language,
            )
        )

    full_text = "\n\n".join(section.text for section in sections).strip()
    if not full_text:
        raise DocumentLoadError("docx document has no text content")

    return ParsedDocument(
        filename=filename,
        file_type="docx",
        char_count=len(full_text),
        preview=full_text[:preview_chars],
        sections=sections,
    )


def _extract_docx_image_ocr_sections(
    *,
    doc: object,
    filename: str,
    start_section_number: int,
    ocr_language: str,
) -> list[ParsedSection]:
    image_parts = [
        part
        for part in doc.part.related_parts.values()
        if str(getattr(part, "content_type", "")).startswith("image/")
    ]

    sections: list[ParsedSection] = []
    for image_index, image_part in enumerate(image_parts, start=1):
        try:
            ocr_image = extract_ocr_text_from_image_bytes(
                filename=filename,
                image_number=image_index,
                content=image_part.blob,
                language=ocr_language,
            )
        except OcrExtractionError as exc:
            raise DocumentLoadError(str(exc)) from exc

        if not ocr_image.text.strip():
            continue

        sections.append(
            ParsedSection(
                section_number=start_section_number + len(sections),
                title=f"docx image {image_index} OCR",
                text=ocr_image.text,
                extraction_method="image_ocr",
            )
        )
    return sections


def _load_csv_document_from_bytes(filename: str, content: bytes, preview_chars: int = 500) -> ParsedDocument:
    text = _decode_text(content)
    rows = list(csv.reader(StringIO(text)))
    if not rows:
        raise DocumentLoadError("csv document has no rows")

    lines = _rows_to_text_lines(rows, sheet_name=None)
    full_text = "\n".join(lines).strip()
    if not full_text:
        raise DocumentLoadError("csv document has no text content")

    return ParsedDocument(
        filename=filename,
        file_type="csv",
        char_count=len(full_text),
        preview=full_text[:preview_chars],
        sections=[ParsedSection(section_number=1, title="csv rows", text=full_text, extraction_method="table")],
    )


def _load_xlsx_document_from_bytes(filename: str, content: bytes, preview_chars: int = 500) -> ParsedDocument:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
    except Exception as exc:
        raise DocumentLoadError(f"Failed to load xlsx document: {exc}") from exc

    sections: list[ParsedSection] = []
    for sheet in workbook.worksheets:
        rows = [
            ["" if value is None else str(value) for value in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        lines = _rows_to_text_lines(rows, sheet_name=sheet.title)
        section_text = "\n".join(lines).strip()
        if section_text:
            sections.append(
                ParsedSection(
                    section_number=len(sections) + 1,
                    title=sheet.title,
                    text=section_text,
                    extraction_method="table",
                )
            )

    if not sections:
        raise DocumentLoadError("xlsx document has no text content")

    full_text = "\n\n".join(section.text for section in sections)
    return ParsedDocument(
        filename=filename,
        file_type="xlsx",
        char_count=len(full_text),
        preview=full_text[:preview_chars],
        sections=sections,
    )


def _load_html_document_from_bytes(filename: str, content: bytes, preview_chars: int = 500) -> ParsedDocument:
    html = _decode_text(content)
    parser = _HtmlBodyTextParser()
    try:
        parser.feed(html)
        parser.close()
    except Exception as exc:
        raise DocumentLoadError(f"Failed to parse html document: {exc}") from exc

    sections: list[ParsedSection] = []
    body_text = parser.body_text()
    if body_text:
        sections.append(
            ParsedSection(
                section_number=len(sections) + 1,
                title=parser.title or "html body",
                text=body_text,
                extraction_method="text",
            )
        )

    table_text = parser.table_text()
    if table_text:
        sections.append(
            ParsedSection(
                section_number=len(sections) + 1,
                title="html tables",
                text=table_text,
                extraction_method="table",
            )
        )

    if not sections:
        raise DocumentLoadError("html document has no body text content")

    full_text = "\n\n".join(section.text for section in sections)
    return ParsedDocument(
        filename=filename,
        file_type="html",
        char_count=len(full_text),
        preview=full_text[:preview_chars],
        sections=sections,
    )


class _HtmlBodyTextParser(HTMLParser):
    _SKIP_TAGS = {"script", "style", "noscript", "svg", "canvas", "template"}
    _NOISE_TAGS = {"nav", "header", "footer", "form", "button"}
    _BLOCK_TAGS = {
        "address",
        "article",
        "aside",
        "blockquote",
        "br",
        "dd",
        "div",
        "dl",
        "dt",
        "figcaption",
        "figure",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "hr",
        "li",
        "main",
        "ol",
        "p",
        "pre",
        "section",
        "tr",
        "ul",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.title: str | None = None
        self._title_parts: list[str] = []
        self._body_parts: list[str] = []
        self._tables: list[list[list[str]]] = []
        self._tag_stack: list[str] = []
        self._skip_depth = 0
        self._noise_depth = 0
        self._in_title = False
        self._current_table: list[list[str]] | None = None
        self._current_row: list[str] | None = None
        self._current_cell_parts: list[str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        self._tag_stack.append(tag)
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        if tag in self._NOISE_TAGS:
            self._noise_depth += 1
        if tag == "title":
            self._in_title = True
        if self._is_skipping:
            return
        if tag in self._BLOCK_TAGS:
            self._append_body_break()
        if tag == "table":
            self._current_table = []
        elif tag == "tr" and self._current_table is not None:
            self._current_row = []
        elif tag in {"td", "th"} and self._current_row is not None:
            self._current_cell_parts = []

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag == "title":
            self._in_title = False
            title = _normalize_html_text(" ".join(self._title_parts))
            self.title = title or self.title

        if not self._is_skipping:
            if tag in {"td", "th"} and self._current_cell_parts is not None and self._current_row is not None:
                cell_text = _normalize_html_text(" ".join(self._current_cell_parts))
                self._current_row.append(cell_text)
                self._current_cell_parts = None
            elif tag == "tr" and self._current_table is not None and self._current_row is not None:
                if any(self._current_row):
                    self._current_table.append(self._current_row)
                self._current_row = None
            elif tag == "table" and self._current_table is not None:
                if self._current_table:
                    self._tables.append(self._current_table)
                self._current_table = None
            if tag in self._BLOCK_TAGS:
                self._append_body_break()

        if tag in self._NOISE_TAGS and self._noise_depth > 0:
            self._noise_depth -= 1
        if tag in self._SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1
        if self._tag_stack:
            self._tag_stack.pop()

    def handle_data(self, data: str) -> None:
        text = _normalize_html_text(data)
        if not text:
            return
        if self._in_title and self._skip_depth == 0:
            self._title_parts.append(text)
            return
        if self._is_skipping:
            return
        if self._current_cell_parts is not None:
            self._current_cell_parts.append(text)
            return
        if self._current_table is None:
            self._body_parts.append(text)

    @property
    def _is_skipping(self) -> bool:
        return self._skip_depth > 0 or self._noise_depth > 0

    def _append_body_break(self) -> None:
        if self._body_parts and self._body_parts[-1] != "\n":
            self._body_parts.append("\n")

    def body_text(self) -> str:
        return _collapse_html_lines(self._body_parts)

    def table_text(self) -> str:
        lines: list[str] = []
        for table_index, table in enumerate(self._tables, start=1):
            lines.extend(_html_table_to_text_lines(table=table, table_index=table_index))
        return "\n".join(lines).strip()


def _html_table_to_text_lines(*, table: list[list[str]], table_index: int) -> list[str]:
    rows = [row for row in table if any(row)]
    if not rows:
        return []
    headers = rows[0]
    has_headers = any(headers)
    lines: list[str] = []
    data_rows = rows[1:] if has_headers else rows
    start_row_number = 2 if has_headers else 1
    for row_index, row in enumerate(data_rows, start=start_row_number):
        values = [str(value).strip() for value in row]
        if not any(values):
            continue
        if has_headers:
            pairs = [
                f"{headers[index] if index < len(headers) and headers[index] else f'column_{index + 1}'}={value}"
                for index, value in enumerate(values)
                if value
            ]
        else:
            pairs = [value for value in values if value]
        if pairs:
            lines.append(f"html table {table_index} row {row_index}: " + "; ".join(pairs))
    return lines


def _collapse_html_lines(parts: list[str]) -> str:
    lines: list[str] = []
    current: list[str] = []
    for part in parts:
        if part == "\n":
            if current:
                lines.append(_normalize_html_text(" ".join(current)))
                current = []
            continue
        current.append(part)
    if current:
        lines.append(_normalize_html_text(" ".join(current)))
    return "\n".join(line for line in lines if line).strip()


def _normalize_html_text(text: str) -> str:
    return " ".join(text.split())


def _rows_to_text_lines(rows: list[list[str]], sheet_name: str | None) -> list[str]:
    if not rows:
        return []

    headers = [str(value).strip() for value in rows[0]]
    has_headers = any(headers)
    lines: list[str] = []
    for row_index, row in enumerate(rows[1:] if has_headers else rows, start=2 if has_headers else 1):
        values = [str(value).strip() for value in row]
        if not any(values):
            continue
        if has_headers:
            pairs = [
                f"{header or f'column_{index + 1}'}={value}"
                for index, (header, value) in enumerate(zip(headers, values))
                if value
            ]
        else:
            pairs = [value for value in values if value]
        prefix = f"sheet {sheet_name} " if sheet_name else ""
        lines.append(f"{prefix}row {row_index}: " + "; ".join(pairs))
    return lines


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentLoadError("Only UTF-8 text files are supported") from exc


def _parse_markdown_sections(text: str) -> list[ParsedSection]:
    sections: list[ParsedSection] = []
    current_title: str | None = None
    current_lines: list[str] = []

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                _append_markdown_section(sections, current_title, current_lines)
                current_title = title
                current_lines = [line]
                continue
        current_lines.append(line)

    _append_markdown_section(sections, current_title, current_lines)
    if not sections:
        sections.append(ParsedSection(section_number=1, title=None, text=text))
    return sections


def _append_markdown_section(
    sections: list[ParsedSection],
    title: str | None,
    lines: list[str],
) -> None:
    section_text = "\n".join(lines).strip()
    if not section_text:
        return
    sections.append(
        ParsedSection(
            section_number=len(sections) + 1,
            title=title,
            text=section_text,
        )
    )
