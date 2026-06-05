import pytest
from docx import Document
from openpyxl import Workbook
from io import BytesIO
from PIL import Image

from app.document_loaders import DocumentLoadError, load_document_from_bytes, load_text_document_from_bytes
from app.ocr_extractor import OcrImage


def test_load_markdown_document_splits_heading_sections():
    document = load_text_document_from_bytes(
        filename="notes.md",
        content="# 标题\n第一段\n\n## 子标题\n第二段".encode("utf-8"),
    )

    assert document.file_type == "markdown"
    assert len(document.sections) == 2
    assert document.sections[0].title == "标题"
    assert "第一段" in document.sections[0].text
    assert document.sections[1].title == "子标题"


def test_load_txt_document_as_single_section():
    document = load_text_document_from_bytes(
        filename="notes.txt",
        content="纯文本知识库内容".encode("utf-8"),
    )

    assert document.file_type == "text"
    assert len(document.sections) == 1
    assert document.sections[0].text == "纯文本知识库内容"


def test_load_text_document_rejects_non_utf8():
    with pytest.raises(DocumentLoadError):
        load_text_document_from_bytes(filename="bad.txt", content=b"\xff\xfe\x00")


def test_load_docx_document_extracts_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("项目代号是 DocxFalcon。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "负责人"
    table.cell(0, 1).text = "Carol"
    table.cell(1, 0).text = "预算"
    table.cell(1, 1).text = "900"
    buffer = BytesIO()
    document.save(buffer)

    parsed = load_document_from_bytes(filename="demo.docx", content=buffer.getvalue())

    assert parsed.file_type == "docx"
    joined_text = "\n".join(section.text for section in parsed.sections)
    assert "DocxFalcon" in joined_text
    assert "Carol" in joined_text
    assert {section.extraction_method for section in parsed.sections} == {"text", "table"}


def test_load_docx_document_extracts_image_ocr(monkeypatch):
    image = Image.new("RGB", (80, 30), color="white")
    image_buffer = BytesIO()
    image.save(image_buffer, format="PNG")
    image_buffer.seek(0)

    document = Document()
    document.add_paragraph("正文内容。")
    document.add_picture(image_buffer)
    docx_buffer = BytesIO()
    document.save(docx_buffer)

    def fake_extract_ocr_text_from_image_bytes(filename, image_number, content, language, preview_chars=500):
        assert filename == "image-demo.docx"
        assert image_number == 1
        assert content.startswith(b"\x89PNG")
        assert language == "eng"
        return OcrImage(
            image_number=image_number,
            char_count=18,
            preview="ImageProject",
            text="ImageProject comes from OCR.",
        )

    monkeypatch.setattr(
        "app.document_loaders.extract_ocr_text_from_image_bytes",
        fake_extract_ocr_text_from_image_bytes,
    )

    parsed = load_document_from_bytes(
        filename="image-demo.docx",
        content=docx_buffer.getvalue(),
        enable_image_ocr=True,
        ocr_language="eng",
    )

    image_sections = [section for section in parsed.sections if section.extraction_method == "image_ocr"]
    assert image_sections
    assert "ImageProject" in image_sections[0].text


def test_load_csv_document_converts_rows_to_text():
    parsed = load_document_from_bytes(
        filename="demo.csv",
        content="name,owner\nCsvRocket,Dana\n".encode("utf-8"),
    )

    assert parsed.file_type == "csv"
    assert "name=CsvRocket" in parsed.sections[0].text
    assert "owner=Dana" in parsed.sections[0].text
    assert parsed.sections[0].extraction_method == "table"


def test_load_xlsx_document_converts_sheets_to_sections():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Projects"
    sheet.append(["name", "owner"])
    sheet.append(["XlsxComet", "Eve"])
    buffer = BytesIO()
    workbook.save(buffer)

    parsed = load_document_from_bytes(filename="demo.xlsx", content=buffer.getvalue())

    assert parsed.file_type == "xlsx"
    assert parsed.sections[0].title == "Projects"
    assert "name=XlsxComet" in parsed.sections[0].text
    assert "owner=Eve" in parsed.sections[0].text
    assert parsed.sections[0].extraction_method == "table"
